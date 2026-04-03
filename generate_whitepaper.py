#!/usr/bin/env python3
"""
generate_whitepaper.py

Generates: AI_PM_Whitepaper_Bosch_MIL.docx

Whitepaper: "Autonomous AI for Project Management — Transforming IT Carve-Out Delivery"
Intended use: Blog post for the official pilot launch (May 2026 onwards).
No confidential engagement details (no project names, buyer/seller identities, or
personnel identifiers) — safe for internal blog and external-facing communication.

Usage:
    "C:/Program Files/px/python.exe" generate_whitepaper.py

Requirements: python-docx  (pip install python-docx)
"""

import sys, os, datetime, base64, io
sys.path.insert(0, os.path.join(os.path.expanduser("~"), "py_packages"))

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

TODAY      = datetime.date.today()
PUB_DATE   = TODAY.strftime("%B %Y")
HERE       = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH  = os.path.join(HERE, "Bosch.png")
OUT_PATH   = os.path.join(HERE, "AI_PM_Whitepaper_Bosch_MIL.docx")

# Bosch colour palette (RGB tuples)
NAVY       = RGBColor(0x00, 0x3b, 0x6e)   # #003b6e  — deep navy
ACCENT     = RGBColor(0x00, 0x66, 0xCC)   # #0066CC  — mid-blue
SECT_BLUE  = RGBColor(0x00, 0x51, 0x99)   # #005199  — section blue
LIGHT_BG   = RGBColor(0xe8, 0xf0, 0xfa)   # #e8f0fa  — light blue tint
MID_GRAY   = RGBColor(0x55, 0x55, 0x55)   # #555555
BODY_TEXT  = RGBColor(0x1a, 0x1a, 0x1a)   # #1a1a1a
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
AMBER      = RGBColor(0xE8, 0xA0, 0x00)   # #E8A000
GREEN      = RGBColor(0x00, 0x7A, 0x33)   # #007A33
RED_BADGE  = RGBColor(0xCC, 0x00, 0x00)   # #CC0000


def set_cell_bg(cell, rgb: RGBColor):
    """Set table cell background shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None,
                     color="003b6e", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val)
            el.set(qn("w:sz"),    size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_run_font(run, size_pt, bold=False, italic=False, color: RGBColor = None,
                 font_name="Calibri"):
    run.font.name      = font_name
    run.font.size      = Pt(size_pt)
    run.font.bold      = bold
    run.font.italic    = italic
    if color:
        run.font.color.rgb = color


def para_space(para, before_pt=0, after_pt=6):
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)


def add_horizontal_rule(doc, color_hex="003b6e", thickness_pt=1):
    """Add a styled paragraph that acts as a horizontal rule."""
    p = doc.add_paragraph()
    para_space(p, 0, 0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(int(thickness_pt * 8)))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def set_table_borders(table, color_hex="003b6e"):
    """Apply uniform single-line borders to entire table."""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_alignment(table, align=WD_TABLE_ALIGNMENT.CENTER):
    tbl   = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    jc    = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_breaks.WD_BREAK.PAGE)


# ─────────────────────────────────────────────────────────────────────────────
# STYLE WRAPPERS
# ─────────────────────────────────────────────────────────────────────────────

def heading1(doc, text, numbered=True, number=""):
    """Section heading — navy bold, large."""
    p   = doc.add_paragraph()
    para_space(p, before_pt=18, after_pt=6)
    if numbered:
        run0 = p.add_run(f"{number}  ")
        set_run_font(run0, 13, bold=True, color=ACCENT)
    run = p.add_run(text.upper())
    set_run_font(run, 13, bold=True, color=NAVY)
    # Bottom border (rule under heading)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003b6e")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    """Sub-section heading — accent blue bold."""
    p   = doc.add_paragraph()
    para_space(p, before_pt=10, after_pt=4)
    run = p.add_run(text)
    set_run_font(run, 11, bold=True, color=SECT_BLUE)
    return p


def body(doc, text, indent_cm=0, justify=True):
    """Standard body paragraph."""
    p   = doc.add_paragraph()
    para_space(p, before_pt=0, after_pt=5)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 10.5, color=BODY_TEXT)
    return p


def body_mixed(doc, parts, indent_cm=0, justify=True):
    """Body paragraph with mixed bold/normal runs.
    parts = list of (text, bold_flag) tuples."""
    p = doc.add_paragraph()
    para_space(p, before_pt=0, after_pt=5)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, 10.5, bold=bold, color=BODY_TEXT)
    return p


def bullet_point(doc, text, level=0):
    """Bullet list item."""
    p   = doc.add_paragraph(style="List Bullet")
    para_space(p, before_pt=1, after_pt=3)
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run(text)
    set_run_font(run, 10.5, color=BODY_TEXT)
    return p


def callout_box(doc, title, lines, bg: RGBColor = None, border_color="0066CC"):
    """Shaded callout / info box using a 1-column table."""
    if bg is None:
        bg = LIGHT_BG
    tbl = doc.add_table(rows=1, cols=1)
    set_table_alignment(tbl)
    set_table_borders(tbl, color_hex=border_color)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Cm(16)

    # Title row
    if title:
        tp = cell.paragraphs[0]
        para_space(tp, 4, 2)
        tr = tp.add_run(title.upper())
        set_run_font(tr, 9, bold=True, color=NAVY)

    for line in lines:
        lp = cell.add_paragraph()
        para_space(lp, 1, 1)
        lp.paragraph_format.left_indent = Cm(0.3)
        bold_part = line.startswith("**")
        clean = line.lstrip("* ")
        lr = lp.add_run(("• " if not bold_part else "") + clean)
        set_run_font(lr, 10, bold=bold_part, color=BODY_TEXT)

    doc.add_paragraph()   # spacer
    return tbl


def add_figure_caption(doc, text):
    p = doc.add_paragraph()
    para_space(p, 2, 8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 9, italic=True, color=MID_GRAY)


# ─────────────────────────────────────────────────────────────────────────────
# HEADER TABLE  (for each page – via section headers would need complex XML;
#               we add it once at the top of the document body instead as a
#               design "letterhead" strip, matching Bosch look)
# ─────────────────────────────────────────────────────────────────────────────

def add_styled_table_header(doc, left_text, right_text,
                             bg: RGBColor = NAVY):
    tbl = doc.add_table(rows=1, cols=2)
    set_table_borders(tbl, color_hex="003b6e")
    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)
    set_cell_bg(lc, bg)
    set_cell_bg(rc, bg)
    lc.width = Cm(12)
    rc.width = Cm(5)

    lp = lc.paragraphs[0]
    para_space(lp, 4, 4)
    lr = lp.add_run(left_text)
    set_run_font(lr, 8, bold=True, color=WHITE)

    rp = rc.paragraphs[0]
    para_space(rp, 4, 4)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(right_text)
    set_run_font(rr, 8, color=RGBColor(0xAA, 0xCC, 0xEE))
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# DATA TABLES
# ─────────────────────────────────────────────────────────────────────────────

def data_table(doc, headers, rows, col_widths_cm,
               note=None, zebra=True):
    """Styled data table with navy header row."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_borders(tbl, color_hex="0066CC")

    # header row
    hdr_row = tbl.rows[0]
    for ci, h_ in enumerate(headers):
        cell = hdr_row.cells[ci]
        set_cell_bg(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        para_space(p, 3, 3)
        r = p.add_run(h_)
        set_run_font(r, 9, bold=True, color=WHITE)

    # data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = RGBColor(0xe8, 0xf0, 0xfa) if (zebra and ri % 2 == 0) else WHITE
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            para_space(p, 2, 2)
            is_first = (ci == 0)
            r = p.add_run(str(cell_text))
            set_run_font(r, 9.5, bold=is_first, color=NAVY if is_first else BODY_TEXT)

    # column widths
    for col_idx, w in enumerate(col_widths_cm):
        for row in tbl.rows:
            row.cells[col_idx].width = Cm(w)

    if note:
        np_ = doc.add_paragraph()
        para_space(np_, 2, 8)
        nr = np_.add_run(f"Note: {note}")
        set_run_font(nr, 8.5, italic=True, color=MID_GRAY)
    else:
        doc.add_paragraph()

    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────

def build_cover(doc):
    # ── Hero band (full-width navy table) ────────────────────────────────────
    hero = doc.add_table(rows=1, cols=2)
    set_table_borders(hero, color_hex="003b6e")
    lc = hero.cell(0, 0)
    rc = hero.cell(0, 1)
    set_cell_bg(lc, NAVY)
    set_cell_bg(rc, NAVY)
    lc.width = Cm(11)
    rc.width = Cm(6)

    # Bosch logo in left cell
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_space(lp, 12, 12)
    if os.path.exists(LOGO_PATH):
        try:
            run_logo = lp.add_run()
            run_logo.add_picture(LOGO_PATH, width=Cm(4))
        except Exception:
            lr = lp.add_run("BOSCH")
            set_run_font(lr, 22, bold=True, color=WHITE)

    # Classification badge in right cell
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_space(rp, 14, 14)
    rr = rp.add_run("INTERNAL  |  BOSCH GROUP")
    set_run_font(rr, 8, bold=True, color=RGBColor(0xAA, 0xCC, 0xEE))

    # ── Accent rule ───────────────────────────────────────────────────────────
    acc = doc.add_table(rows=1, cols=1)
    set_table_borders(acc, color_hex="0066CC")
    set_cell_bg(acc.cell(0, 0), ACCENT)
    acc.cell(0, 0).width = Cm(17)
    sp = acc.cell(0, 0).paragraphs[0]
    para_space(sp, 1, 1)

    # ── Title block ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    t1 = doc.add_paragraph()
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(t1, 20, 4)
    r1 = t1.add_run("AUTONOMOUS AI FOR PROJECT MANAGEMENT")
    set_run_font(r1, 22, bold=True, color=NAVY, font_name="Calibri")

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(t2, 4, 4)
    r2 = t2.add_run("Transforming IT Carve-Out Delivery")
    set_run_font(r2, 16, bold=False, color=SECT_BLUE, font_name="Calibri")

    doc.add_paragraph()

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(t3, 4, 4)
    r3 = t3.add_run(
        "How AI-augmented domain knowledge is eliminating the documentation burden "
        "for IT carve-out project managers — and what this means for the future of M&A delivery"
    )
    set_run_font(r3, 11, italic=True, color=MID_GRAY, font_name="Calibri")

    # ── Divider ───────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()
    add_horizontal_rule(doc, color_hex="0066CC", thickness_pt=1.5)

    # ── Meta strip table ──────────────────────────────────────────────────────
    meta = doc.add_table(rows=4, cols=2)
    set_table_borders(meta, color_hex="CCDDEE")
    meta_data = [
        ("Practice",        "Mergers, Integrations & Liquidations (MIL)"),
        ("Organisation",    "Bosch Global Business Services — IT Strategy"),
        ("Published",       PUB_DATE),
        ("Classification",  "Internal — Bosch Group"),
    ]
    for ri, (label, val) in enumerate(meta_data):
        lc_ = meta.rows[ri].cells[0]
        rc_ = meta.rows[ri].cells[1]
        set_cell_bg(lc_, RGBColor(0xf4, 0xf6, 0xf9))
        set_cell_bg(rc_, WHITE)
        lc_.width = Cm(4)
        rc_.width = Cm(13)
        lp_ = lc_.paragraphs[0]
        para_space(lp_, 3, 3)
        lr_ = lp_.add_run(label)
        set_run_font(lr_, 9, bold=True, color=NAVY)
        rp_ = rc_.paragraphs[0]
        para_space(rp_, 3, 3)
        rr_ = rp_.add_run(val)
        set_run_font(rr_, 9, color=BODY_TEXT)

    for _ in range(4):
        doc.add_paragraph()

    # ── Abstract callout ──────────────────────────────────────────────────────
    callout_box(doc, "Executive Summary", [
        "IT carve-out projects generate enormous documentation burdens before a single line of code "
        "is written. Schedules, risk registers, cost plans, project charters, executive dashboards, "
        "KPI dashboards, and monthly reports must all be produced — and kept consistent with each other — "
        "from day one of mobilisation.",
        "This whitepaper describes a pilot AI solution developed within Bosch's MIL practice that "
        "automates the generation of these deliverables using GitHub Copilot (Claude Sonnet AI) integrated "
        "with deep M&A methodology knowledge. The system combines VS Code as the workspace, GitHub as the "
        "knowledge repository, and a structured layer of domain-specific skills to produce "
        "governance-quality documents in minutes — not days.",
        "The pilot programme launches in May 2026. This paper introduces the architecture, the value "
        "proposition, and the path forward.",
    ])

    # Page break before body
    p = doc.add_paragraph()
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    pPr = p._p.get_or_add_pPr()
    pb = _OE("w:pageBreakBefore")
    pb.set(_qn("w:val"), "true")
    pPr.append(pb)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION CONTENT
# ─────────────────────────────────────────────────────────────────────────────

def section_problem(doc):
    heading1(doc, "The Challenge: Documentation Overload at Project Start", number="01")

    body(doc,
        "Every IT carve-out project begins the same way. A project manager is assigned, "
        "a timeline is set, and within days the requests arrive: a schedule for the Steering "
        "Committee, a risk register for legal review, a cost plan for finance approval, a "
        "project charter for the executive sponsor, and a dashboard for the monthly reporting "
        "cycle. All before the team has held its first discovery workshop.")

    body(doc,
        "This documentation burden is not administrative overhead — it is foundational project "
        "governance. The schedule sets the baseline; the risk register informs the budget; "
        "the cost plan supports investment approval; the charter secures stakeholder alignment. "
        "Get any one of them wrong, or leave them inconsistent with each other, and the "
        "consequences surface at the worst possible moment: a Quality Gate review, a SteerCo "
        "escalation, or a due-diligence challenge from the buyer's legal team.")

    heading2(doc, "The Real Cost of Manual Document Production")

    body(doc,
        "For a typical mid-size carve-out — spanning multiple sites, several hundred to thousands "
        "of users, and a 9–18 month timeline — producing the initial document set manually takes "
        "an experienced project manager between three and five working days. That is before any "
        "review, revision, or stakeholder alignment. And it must be repeated, in part, at every "
        "monthly reporting cycle.")

    data_table(doc,
        headers=["Deliverable", "Manual Effort", "Key Risk if Inconsistent"],
        rows=[
            ("Project Schedule (CSV + MS Project XML)",
             "5–7 days",
             "Complex phase/task dependencies require many iterations; GoLive date easily misaligned"),
            ("Risk Register (.xlsx)",
             "3 days",
             "Uncosted mitigations; scoring inconsistent without structured methodology"),
            ("IT Labour Cost Plan (.csv)",
             "3 days",
             "Phase costs not tied to schedule; resource labels drift between documents"),
            ("Project Charter (.html / .pptx)",
             "3 days",
             "Scope definition differs from schedule; executive misalignment at QG0"),
            ("Executive Dashboard (currently PPT slides)",
             "5+ days",
             "HTML-based reporting beyond most PMs — PPT alternative is labour-intensive and not self-updating"),
            ("Management KPI Dashboard (currently PPT slides)",
             "5+ days",
             "Same challenge as Executive Dashboard; SPI/CPI metrics require manual recalculation each cycle"),
            ("Monthly Status Report (.pdf)",
             "2–3 days/month",
             "Manual date recalculations each month; phase RAG requires cross-referencing multiple source files"),
        ],
        col_widths_cm=[5.5, 2.7, 8.8],
        note="Estimates based on real MIL engagement experience across multiple carve-out programmes. Excludes review and approval cycles.",
    )

    body(doc,
        "Beyond time, there is a consistency problem. Each document is typically produced in "
        "isolation — a different tool, a different template, a different revision of the schedule. "
        "The risk register references phase names from an early draft of the schedule. The cost "
        "plan uses resource labels that were later renamed. The dashboard shows milestones from "
        "last month's version. These gaps are not negligence — they are the inevitable result of "
        "fragmented, manual, time-pressured production.")


def section_solution(doc):
    heading1(doc, "The Solution: AI-Assisted Document Generation", number="02")

    body(doc,
        "The Bosch MIL AI Project Management Toolkit addresses this challenge by combining two "
        "things that were previously separate: the general reasoning capability of large language "
        "models, and the deep, specific methodological knowledge accumulated across years of "
        "Bosch IT carve-out delivery.")

    body(doc,
        "The result is not a chatbot that writes documents when asked. It is an orchestrated "
        "AI system that knows the Bosch MIL methodology, enforces it, applies it to the current "
        "engagement's specific facts, cross-checks its own outputs for consistency, and produces "
        "complete, governance-quality deliverables — in minutes, from a single conversation.")

    callout_box(doc,
        "What makes this different from 'asking AI to write a document'",
        [
            "The AI reads structured methodology instructions before any generation begins.",
            "Mandatory intake validation blocks generation if any of 11 required fields are missing.",
            "Each deliverable is generated in a dependency-enforced sequence — no skipping steps.",
            "Cross-checks between documents are non-negotiable: resource names, phase dates, "
            "and cost contingencies are verified for consistency before output is produced.",
            "All outputs — Python scripts, HTML, CSV, XML, PDF — are committed to Git, creating "
            "a full audit trail of what was generated, when, and under what instructions.",
        ],
        border_color="0066CC",
    )

    heading2(doc, "From Prompt to Governance-Ready Deliverable")

    body(doc,
        "The workflow begins with a single conversation. The project manager opens "
        "GitHub Copilot in VS Code Agent mode and provides the eleven mandatory engagement "
        "parameters — project name, seller entity, buyer entity, business being carved out, "
        "carve-out model, PMO lead, worldwide site count, IT user count, start date, GoLive "
        "date, and completion date. The AI validates these, derives all dependent fields, and "
        "then generates the full deliverable set in the correct sequence.")

    data_table(doc,
        headers=["Input", "AI Action", "Output"],
        rows=[
            ("11 engagement parameters",
             "Intake compliance gate — validates all fields; blocks if any missing",
             "Validation confirmation + derived buyer/seller governance context"),
            ("Confirmed engagement context",
             "Applies Bosch 5-phase / 5-QG framework; populates task list",
             "Project Schedule CSV + MSPDI XML (MS Project-ready)"),
            ("Schedule confirmed",
             "Applies carve-out risk taxonomy (8 categories, P×I scoring)",
             "Risk Register .xlsx (template-based, formula-ready)"),
            ("Schedule + Risk Register confirmed",
             "Cross-checks phases, resources, contingency lines",
             "IT Labour Cost Plan CSV with phase breakdown"),
            ("Cost Plan confirmed",
             "Composes governance narrative; embeds Bosch logo",
             "Project Charter self-contained HTML"),
            ("Full data set confirmed",
             "Renders 3-page A4 layout; base64-embeds logo for offline use",
             "Executive Dashboard self-contained HTML"),
            ("Full data set confirmed",
             "12-column KPI card grid; SPI/CPI/readiness metrics",
             "Management KPI Dashboard self-contained HTML"),
            ("All above confirmed",
             "A4 PDF with auto-calculated days-to-gate from today's date",
             "Monthly Status Report PDF (auto-dated filename)"),
        ],
        col_widths_cm=[4.5, 6.5, 6],
    )


def section_architecture(doc):
    heading1(doc, "The Architecture: How the System is Built", number="03")

    body(doc,
        "The toolkit is built on four integrated layers. Understanding each layer — and how they "
        "interact — explains why this approach produces consistently correct output, rather than "
        "the plausible-but-wrong output that emerges from simple AI prompting.")

    heading2(doc, "Layer 1 — The AI Engine: GitHub Copilot on Claude Sonnet")

    body(doc,
        "GitHub Copilot, powered by Anthropic's Claude Sonnet 4.6 model, serves as the AI reasoning "
        "engine. Running inside Visual Studio Code in Agent mode, it can read files, write files, "
        "execute Python scripts, and take multi-step actions — not just generate text responses. "
        "This agentic capability is what makes multi-deliverable orchestration possible in a "
        "single conversation.")

    body_mixed(doc, [
        ("Agent mode is critical. ", True),
        ("In standard chat mode, Copilot suggests; in Agent mode, Copilot acts. It reads the "
         "methodology skill files, writes Python generator scripts, executes them, reads the "
         "output, cross-checks for consistency, and reports completion — all within one workflow.", False),
    ])

    heading2(doc, "Layer 2 — The Knowledge Repository: GitHub + VS Code Workspace")

    body(doc,
        "The knowledge repository on GitHub is the institutional brain of the system. It stores "
        "not just code and templates, but the methodology instructions that tell the AI "
        "how to behave. Two files are loaded automatically by Copilot every time the workspace "
        "is opened:")

    bullet_point(doc,
        "copilot-instructions.md (.github/): Workspace-level guardrails — the always-on "
        "constitution that enforces intake validation, derivation rules, output format standards, "
        "and the mandatory deliverable sequence.")
    bullet_point(doc,
        "CLAUDE.md (repository root): A parallel copy of the same guardrails in Claude's native "
        "format — ensuring the rules apply regardless of which AI entry point is used.")

    body(doc,
        "The repository follows a four-zone content model: references/ (methodology assets, "
        "read-only to AI), templates/ (blank project templates), active-projects/ (live "
        "engagement outputs), and archive/ (closed engagements). This structure prevents "
        "a common AI failure mode — copying facts from a reference project into a new engagement.")

    heading2(doc, "Layer 3 — The Domain Skills: Encoded MIL Methodology")

    body(doc,
        "The most important architectural element is the skills layer. Eight structured "
        "SKILL.md files in the .claude/skills/ directory encode the complete Bosch MIL "
        "methodology for each deliverable type. These are not prompt templates — they are "
        "procedural specifications that the AI reads and executes.")

    data_table(doc,
        headers=["Skill", "What It Encodes"],
        rows=[
            ("intake-compliance-gate",
             "11 mandatory fields, blocking rules, budget TBC exception, buyer/seller derivation"),
            ("schedule-generation",
             "5-phase/5-QG framework, task schema, MSPDI XML element ordering rules, subprocess pattern"),
            ("cost-plan-generation",
             "3 pre-generation cross-checks: schedule alignment, risk register alignment, resource consistency"),
            ("risk-register-generation",
             "Risk taxonomy (8 categories), P×I scoring, template column mapping, openpyxl rules"),
            ("executive-dashboard-generation",
             "3-page A4 Trinity layout, Bosch colour palette, logo embedding, self-contained HTML rules"),
            ("management-kpi-dashboard-generation",
             "12-column card grid, SPI/CPI/readiness metrics, 90-day action forecast"),
            ("monthly-status-report-generation",
             "A4 reportlab layout, auto-dated filename, days-to-gate auto-calculation"),
            ("repository-governance-updates",
             "Metadata maintenance rules, naming conventions, inventory sync, last-reviewed tracking"),
        ],
        col_widths_cm=[5.5, 11.5],
    )

    body(doc,
        "When the AI receives a generation request, it looks up the relevant skill in the routing "
        "table defined in copilot-instructions.md, reads the full SKILL.md file, and executes "
        "its instructions. This two-level architecture keeps the global instruction file "
        "concise — fitting within the AI's context window — while allowing each skill to carry "
        "rich, detailed methodology logic.")

    heading2(doc, "Layer 4 — The Generator Scripts: Python Automation")

    body(doc,
        "The AI does not write documents by filling in Word templates. It writes Python scripts "
        "that generate documents programmatically — using reportlab for PDFs, openpyxl for Excel "
        "files, and pure Python for MSPDI XML and self-contained HTML. Each project has its own "
        "set of generator scripts that hardcode the engagement parameters as auditable Python "
        "data structures.")

    body(doc,
        "The canonical CSV-to-XML converter (generate_msp_xml.py) is a shared script called "
        "by all schedule generators as a subprocess. It encodes a set of critical Microsoft "
        "Project MSPDI XML rules discovered empirically through real import testing — rules "
        "that are not documented in the MS Project specification but which are essential for "
        "task dates to survive import without drifting. These rules are captured in the skill "
        "file, the script, and the AI's persistent memory layer.")

    callout_box(doc,
        "The MS Project XML date integrity problem — solved",
        [
            "Problem: MS Project recalculates all task dates from the predecessor chain on import, "
            "causing GoLive to appear months early or late.",
            "Cause: The <TaskMode>1</TaskMode> element must appear immediately after <Name> in each "
            "Task element. Placed anywhere else, MS Project silently ignores it.",
            "Solution: generate_msp_xml.py enforces strict element ordering, ManualStart/Finish "
            "fields, ConstraintType='Must Start On', and project-level manual scheduling flags.",
            "Impact: All project schedules import into MS Project with dates exactly as authored — "
            "zero drift, zero manual correction required.",
        ],
        bg=RGBColor(0xef, 0xf5, 0xff),
        border_color="003b6e",
    )

    heading2(doc, "The Memory System: AI That Learns Across Sessions")

    body(doc,
        "A key differentiator of this system is its three-tier persistent memory architecture. "
        "Unlike a standard AI session that forgets everything when closed, this system maintains "
        "structured memory files that survive across all sessions:")

    bullet_point(doc,
        "User memory: Cross-workspace permanent notes capturing output standards, known bugs "
        "and their fixes, and CSS/XML rules that must not be regressed.")
    bullet_point(doc,
        "Session memory: Conversation-scoped working notes for multi-step tasks.")
    bullet_point(doc,
        "Repository memory: Repository-scoped facts — GitHub remote URL, Python interpreter "
        "path, canonical branch.")

    body(doc,
        "When a new session starts, the AI reads these memory files before acting. It does not "
        "rediscover known issues. It does not reproduce previously fixed bugs. The system "
        "becomes more reliable with every engagement — accumulating institutional knowledge "
        "in a structured, machine-readable form.")


def section_value(doc):
    heading1(doc, "Value Proposition: What This Changes for Project Managers", number="04")

    body(doc,
        "The value of this system is not simply speed — though the speed is dramatic. It is "
        "a fundamental shift in how project managers can deploy their expertise at the start "
        "of a carve-out engagement.")

    heading2(doc, "From Weeks to Under One Hour: Speed of Mobilisation")

    body(doc,
        "A complete initial deliverable set that previously required three to four weeks of "
        "expert manual work is generated in under 35 minutes of AI-assisted production. "
        "For context: the project schedule alone — with all its phase dependencies, QG milestones, "
        "and resource assignments — takes an experienced PM five to seven days to construct from "
        "scratch manually. The two dashboards are even more challenging: most PMs are not fluent in "
        "HTML-based reporting and rely on PowerPoint slides instead, each taking more than five days "
        "to produce per reporting cycle. The AI generates all seven deliverables in a single "
        "35-minute automated run — each item taking approximately five minutes.")

    callout_box(doc,
        "Time comparison: manual vs AI-assisted production",
        [
            "Schedule (CSV + MS Project XML)       Manual: 5–7 days    →  AI: ~5 minutes",
            "Risk Register (.xlsx, 15–25 risks)    Manual: 3 days      →  AI: ~5 minutes",
            "IT Labour Cost Plan (.csv)             Manual: 3 days      →  AI: ~5 minutes",
            "Project Charter (HTML)                 Manual: 3 days      →  AI: ~5 minutes",
            "Executive Dashboard (HTML — vs PPT)    Manual: 5+ days     →  AI: ~5 minutes",
            "Management KPI Dashboard (HTML — vs PPT) Manual: 5+ days   →  AI: ~5 minutes",
            "Monthly Status Report (PDF)            Manual: 2–3 days/month → AI: ~5 minutes",
            "TOTAL initial set:  3–4 weeks  →  Under 35 minutes (full automated run)",
        ],
        bg=RGBColor(0xe6, 0xf4, 0xec),
        border_color="007A33",
    )

    heading2(doc, "Structural Consistency: Documents That Actually Agree")

    body(doc,
        "Because all deliverables are generated from the same source facts — the engagement "
        "parameters supplied at intake — and because the AI enforces cross-checks between them, "
        "the outputs are structurally consistent in a way that manual production rarely achieves:")

    bullet_point(doc,
        "Phase names and dates in the cost plan match the schedule exactly — because the "
        "cost plan skill mandates a schedule alignment cross-check before writing any figures.")
    bullet_point(doc,
        "Resource labels in the cost plan map to resource names in the schedule — because "
        "the skill checks every resource token for traceability.")
    bullet_point(doc,
        "CAPEX contingency lines in the cost plan reference the specific risk register entries "
        "that drove them — because the skill requires a risk register cross-check for all "
        "Amber/Red or high-scoring risks.")
    bullet_point(doc,
        "The executive dashboard countdown to GoLive matches the GoLive date in the schedule — "
        "because both are derived from the same intake parameter.")

    heading2(doc, "Methodological Guardrails: Enforcing What Experience Learned")

    body(doc,
        "The skills layer encodes hard-won methodological lessons that might otherwise be "
        "silently omitted under time pressure. The intake compliance gate will not let a "
        "document be generated if site count, user count, or project dates are missing — "
        "because these are not cosmetic details; they are the parameters that drive schedule "
        "duration, resource loading, and risk exposure.")

    body(doc,
        "Bosch's standard 5-phase, 5-Quality-Gate framework is applied consistently across "
        "every engagement, regardless of size or model (Stand Alone, Integration, or "
        "Combination). The framework is not a suggestion — it is the authoritative structure "
        "that Bosch Quality Management and Legal require at every gate review.")

    heading2(doc, "Audit Trail and Governance")

    body(doc,
        "Every generated file is committed to the Git repository. Every script execution "
        "is timestamped in the terminal output. This creates an immutable, reviewable record "
        "of what was generated, when, under what instructions, and with what input parameters. "
        "For M&A transactions — which involve legal, regulatory, and fiduciary obligations — "
        "this level of traceability is not a nice-to-have. It is essential.")

    heading2(doc, "Focus Where It Matters: Judgment Over Administration")

    body(doc,
        "The most profound benefit is the redirection of the project manager's attention. "
        "When a PM no longer spends three to four weeks building documents, they spend that time on "
        "the questions that actually require human expertise: Which risks are the ones that "
        "will derail this programme? Is the GoLive date achievable given what we know about "
        "the buyer's IT readiness? How do we message the budget ask to the Steering "
        "Committee without triggering a scope reduction? These are judgment calls. "
        "Document production is not.")

    data_table(doc,
        headers=["Without AI Assistance", "With AI Assistance"],
        rows=[
            ("3–4 weeks building the initial document set",
             "Under 35 minutes — PM reviews and approves"),
            ("Documents produced in isolation per tool",
             "All documents generated from a single consistent data set"),
            ("Manual consistency checking across files",
             "Cross-checks enforced automatically before generation"),
            ("Methodology applied from memory, under time pressure",
             "Methodology encoded in skills — always applied, always current"),
            ("No formal link between risk register and cost plan",
             "Every high-scoring risk traces to a cost contingency line"),
            ("Monthly report requires manual re-entry of dates",
             "All dates auto-calculated from today at script run time"),
            ("Audit trail is 'email thread + shared folder'",
             "Full Git commit history — who, what, when, under what instruction"),
        ],
        col_widths_cm=[8.5, 8.5],
    )


def section_deliverables(doc):
    heading1(doc, "The Deliverables: What Gets Generated", number="05")

    body(doc,
        "The system generates seven mandatory deliverables, in a fixed dependency sequence. "
        "No deliverable can be generated before its predecessors are confirmed — this "
        "enforcement is built into the AI instruction layer, not left to the PM's discretion.")

    heading2(doc, "1. Project Schedule — CSV + Microsoft Project XML")

    body(doc,
        "The foundation of all downstream deliverables. Built on Bosch's standard 5-phase "
        "framework (Initialization → Concept → Development & Build → Testing & Cutover → "
        "GoLive & Hypercare) with QG0 through QG5 milestone gates. Delivered in both human-"
        "readable CSV and MSPDI XML format — importable directly into Microsoft Project with "
        "all task dates preserved exactly as authored.")

    heading2(doc, "2. Risk Register — Excel (.xlsx)")

    body(doc,
        "Generated from the authoritative Risk_analysis_template.xlsx with all 8 Bosch risk "
        "categories (Schedule, Scope, Resource, Budget, Quality, Business, Legal, Compliance), "
        "Probability × Impact scoring (1–5 scale), mitigation actions, owners, and target "
        "dates. High-priority risks (P×I ≥ 12) are flagged automatically.")

    heading2(doc, "3. IT Labour Cost Plan — CSV")

    body(doc,
        "A structured cost breakdown derived from schedule resource assignments, with category "
        "subtotals, phase breakdown, and CAPEX/contingency lines cross-referenced to the risk "
        "register. All three components — schedule, risk register, and resource names — are "
        "verified for consistency before any cost figure is written.")

    heading2(doc, "4. Project Charter — Self-Contained HTML")

    body(doc,
        "A governance document presenting scope, timeline, budget baseline, quality gate "
        "criteria, key parties, and programme assumptions. Includes embedded Bosch logo. "
        "Self-contained — no external links — so it renders correctly in SharePoint, "
        "email clients, and offline environments.")

    heading2(doc, "5. Executive Dashboard — Self-Contained HTML (3-Page A4)")

    body(doc,
        "A strategic programme overview including: countdown to key events, phase timeline, "
        "milestones and QG table, budget distribution, workstream confidence grid (9 "
        "workstreams), quality gate tracker, regional site distribution, risk indicators, "
        "application migration waves, country complexity hotspots, and critical path. "
        "Blue-primary Bosch theme. Print-ready — PDF output from browser in one click.")

    heading2(doc, "6. Management KPI Dashboard — Self-Contained HTML")

    body(doc,
        "The operational steering view for PMO and SteerCo. Twelve-column card grid with: "
        "SPI, CPI, Day-1 readiness score, TSA confidence, workstream confidence bars, "
        "milestone gate control timeline, top risk table, carve-out model key differences, "
        "and 90-day action forecast.")

    heading2(doc, "7. Monthly Status Report — PDF")

    body(doc,
        "A single-page A4 report with programme overview, days-to-gate countdown, phase "
        "RAG statuses, risk summary, budget burn, and upcoming actions. Output filename "
        "auto-includes the current month and year — running the script on the first of each "
        "month produces a correctly-named report with all date calculations refreshed.")


def section_pilot(doc):
    heading1(doc, "The Pilot Programme: May 2026 Launch", number="06")

    body(doc,
        "The AI-Assisted Project Management Toolkit has been developed and validated through "
        "a series of internal carve-out engagements, running in parallel as a shadow delivery "
        "track alongside the conventional approach. The pilot produced complete, consistent "
        "deliverable sets that passed internal methodology review — in multiple different "
        "carve-out contexts ranging from small JV separations to large multi-site "
        "Stand Alone programmes.")

    callout_box(doc,
        "Pilot programme — key facts",
        [
            "**Launch month:  May 2026",
            "**Scope:  Active MIL IT carve-out engagements nominated for the pilot",
            "**Model:  AI-generated deliverables as the primary production method, "
            "reviewed and approved by the PM before submission",
            "**Governance:  All outputs committed to Bosch DevCloud Git repository with "
            "full timestamped audit trail",
            "**Platform:  GitHub Copilot (Claude Sonnet 4.6) · VS Code Agent mode · "
            "Python 3 generator scripts",
            "**Success criteria:  Document accuracy, consistency, PM effort reduction, "
            "stakeholder acceptance at QG0 / QG1",
        ],
        border_color="0066CC",
    )

    heading2(doc, "How the Pilot Works in Practice")

    body(doc,
        "A project manager opens the workspace in VS Code and invokes Copilot in Agent mode "
        "with a simple request: 'Generate the full initial deliverable set for this engagement.' "
        "The AI applies the compliance gate, asks for any missing parameters, generates each "
        "deliverable in sequence, performs all cross-checks, and reports completion with "
        "file paths and generation times. The PM reviews the outputs, makes any judgment-level "
        "adjustments to tone or narrative, and commits the approved set to the repository.")

    body(doc,
        "For monthly reporting, the PM runs the status report generator script — a single "
        "command — and receives a correctly dated, current-data PDF ready for distribution. "
        "All days-to-gate values, phase status calculations, and risk summaries update "
        "automatically from today's date.")

    heading2(doc, "What We Are Learning")

    bullet_point(doc,
        "Speed of adoption: How quickly can a PM with no prior AI tooling experience produce "
        "a complete deliverable set unassisted?")
    bullet_point(doc,
        "Review accuracy: What proportion of AI-generated content requires substantive "
        "correction vs. minor factual adjustment?")
    bullet_point(doc,
        "Stakeholder acceptance: Do QG0 and QG1 reviewers accept AI-generated documentation "
        "as methodologically sound?")
    bullet_point(doc,
        "Consistency over time: Does the AI maintain cross-document consistency across "
        "multiple revision cycles during the engagement?")
    bullet_point(doc,
        "Memory effectiveness: Does the three-tier memory system successfully prevent "
        "regression to previously resolved issues?")


def section_future(doc):
    heading1(doc, "Looking Ahead: The Evolving AI Delivery Platform", number="07")

    body(doc,
        "The pilot is the first step in a broader vision: an AI delivery platform for "
        "Bosch MIL that continuously improves with each engagement, scales across the "
        "practice, and eventually supports the full delivery lifecycle — not just the "
        "mobilisation phase.")

    data_table(doc,
        headers=["Capability", "Description", "Timeline"],
        rows=[
            ("Dashboard auto-refresh",
             "Scripts that read live progress data and regenerate "
             "dashboards automatically for weekly status cycles",
             "Q3 2026"),
            ("Deliverable diff reporting",
             "AI comparison of two schedule or cost plan versions, "
             "generating a change-control summary for SteerCo",
             "Q3 2026"),
            ("Portfolio view",
             "Aggregated multi-engagement dashboard showing status "
             "across all active MIL carve-outs from a single view",
             "Q4 2026"),
            ("Natural language status updates",
             "PM types a status narrative; AI propagates updates to "
             "all relevant fields in report and dashboards",
             "Q4 2026"),
            ("Quality gate checklist automation",
             "AI generates gate-readiness checklists from the schedule "
             "QG milestones and open high-priority risk register items",
             "Q1 2027"),
            ("Template versioning",
             "Formal versioning of all templates and layout specs; "
             "AI instructed to use the version tagged at engagement start",
             "Q1 2027"),
        ],
        col_widths_cm=[4.5, 9, 3.5],
    )

    body(doc,
        "The longer-term vision is a practice where new MIL engagements are mobilised in "
        "hours, not days — where the AI handles the systematic application of methodology, "
        "and PMs and consultants focus entirely on the complex, judgement-intensive work "
        "that defines the value of experienced M&A practitioners.")


def section_conclusion(doc):
    heading1(doc, "Conclusion", number="08")

    body(doc,
        "The Bosch MIL AI Project Management Toolkit is not a technology experiment. It is "
        "a practical solution to a real and recurring problem: the documentation burden that "
        "consumes expert time at the most critical phase of every IT carve-out engagement.")

    body(doc,
        "By combining GitHub Copilot's AI reasoning with structured MIL methodology knowledge "
        "— encoded in skill files, enforced by guardrails, and accumulated in persistent memory "
        "— the system produces governance-quality deliverables that are faster to generate, "
        "more consistent across documents, more methodologically rigorous, and more auditable "
        "than their manually-produced equivalents.")

    body(doc,
        "The pilot launching in May 2026 will test this in live engagement conditions. "
        "The architecture is built, the tools are validated, the methodology is encoded. "
        "What remains is to demonstrate — at scale, in real programmes — that autonomous "
        "AI can be a trusted, reliable, and transparent partner in M&A project delivery.")

    body_mixed(doc, [
        ("The construction is automated. The judgment remains human. ", True),
        ("That is the right balance for where AI capability stands today — and it is the "
         "balance this toolkit is designed to preserve.", False),
    ])

    doc.add_paragraph()
    add_horizontal_rule(doc, color_hex="0066CC")
    doc.add_paragraph()

    # Closing meta strip
    end_tbl = doc.add_table(rows=1, cols=3)
    set_table_borders(end_tbl, color_hex="CCDDEE")
    labels = [
        ("Practice",    "Bosch MIL — IT Strategy"),
        ("Platform",    "GitHub Copilot · VS Code · Claude Sonnet"),
        ("Pilot Launch","May 2026"),
    ]
    widths = [5, 7, 5]
    for ci, (label, val) in enumerate(labels):
        cell = end_tbl.rows[0].cells[ci]
        set_cell_bg(cell, RGBColor(0xf4, 0xf6, 0xf9))
        cell.width = Cm(widths[ci])
        p = cell.paragraphs[0]
        para_space(p, 4, 4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "\n")
        set_run_font(r1, 8, bold=True, color=NAVY)
        r2 = p.add_run(val)
        set_run_font(r2, 8, color=BODY_TEXT)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(foot, 4, 0)
    fr = foot.add_run(
        f"© Bosch Group  ·  Internal  ·  Not for external distribution  ·  {PUB_DATE}"
    )
    set_run_font(fr, 8, italic=True, color=MID_GRAY)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_TEXT

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    # Lazy import to avoid error at module level if missing
    from docx.oxml import OxmlElement  # noqa (already imported above)
    try:
        from docx.enum.text import WD_BREAK as _b
        global docx_breaks
        import types
        docx_breaks = types.SimpleNamespace(WD_BREAK=_b)
    except Exception:
        pass

    doc = setup_document()

    # ── Cover ─────────────────────────────────────────────────────────────────
    build_cover(doc)

    # ── Running header (strip at top of each body page — via table, not true header)
    add_styled_table_header(
        doc,
        left_text="AUTONOMOUS AI FOR PROJECT MANAGEMENT  |  Bosch MIL  |  Whitepaper",
        right_text=PUB_DATE,
    )
    add_horizontal_rule(doc, color_hex="003b6e")
    doc.add_paragraph()

    # ── Sections ──────────────────────────────────────────────────────────────
    section_problem(doc)
    section_solution(doc)
    section_architecture(doc)
    section_value(doc)
    section_deliverables(doc)
    section_pilot(doc)
    section_future(doc)
    section_conclusion(doc)

    doc.save(OUT_PATH)
    print(f"Word document written: {OUT_PATH}")
    size_kb = os.path.getsize(OUT_PATH) / 1024
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    from datetime import datetime as _dt
    _t0 = _dt.now()
    print(f"Started : {_t0.strftime('%Y-%m-%d %H:%M:%S')}")
    try:
        main()
    finally:
        _t1 = _dt.now()
        print(f"Finished: {_t1.strftime('%Y-%m-%d %H:%M:%S')}  ({(_t1-_t0).total_seconds():.1f}s elapsed)")
